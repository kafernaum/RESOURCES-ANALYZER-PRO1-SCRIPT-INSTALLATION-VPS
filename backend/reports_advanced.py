"""Advanced report generation:
- REJD complete (8 parts + 8 annexes) extending base reports.py
- Word (.docx) editable export via python-docx
- Excel (.xlsx) multi-sheet via openpyxl
- ZIP pack
"""
import io
import json
import zipfile
from datetime import datetime
from typing import Dict, Any
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side


PRIMARY_HEX = "1B4332"
SECONDARY_HEX = "D4A017"
TERTIARY_HEX = "1A3C5E"
ALERT_HEX = "C0392B"
LIGHT_BG_HEX = "F5F9F5"


# ============== WORD ==============
def _set_cell_bg(cell, hex_color: str):
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    cell._tc.get_or_add_tcPr().append(
        parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), hex_color))
    )


def generate_word(project: Dict[str, Any], report_data: Dict[str, Any], preset: str) -> bytes:
    doc = DocxDocument()
    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAPPORT D'EXPERTISE\n")
    run.bold = True; run.font.size = Pt(28); run.font.color.rgb = RGBColor(0x1B, 0x43, 0x32)
    run2 = title.add_run("RESOURCES-ANALYZER PRO\n")
    run2.bold = True; run2.font.size = Pt(20); run2.font.color.rgb = RGBColor(0xD4, 0xA0, 0x17)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"\nConvention d'exploitation — {project.get('sector', '').capitalize()}")
    sr.italic = True; sr.font.size = Pt(13)
    sub.add_run(f"\n{project.get('country', '')} — {project.get('resource_type', '')}\n")

    doc.add_paragraph(f"\n\nDate d'analyse : {datetime.now().strftime('%d %B %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = p.add_run("\nMéthodologie : Ahmed ELY Mustapha — Juriste, Expert en Finances Publiques")
    pr.italic = True

    discl = doc.add_paragraph()
    discl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = discl.add_run("\nDOCUMENT PÉDAGOGIQUE — n'a aucune valeur juridictionnelle.")
    dr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B); dr.italic = True

    doc.add_page_break()

    # Synthesis section
    summary = report_data.get("summary", {})
    counters = summary.get("compteurs", {})
    h1 = doc.add_heading("1. Synthèse exécutive", level=1)
    for run_h in h1.runs:
        run_h.font.color.rgb = RGBColor(0x1B, 0x43, 0x32)

    table = doc.add_table(rows=6, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["Indicateur", "Valeur", "Niveau"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True

    rows = [
        ["Score Conformité Juridique", f"{summary.get('score_juridique', 0)} / 100", _level(summary.get('score_juridique', 0))],
        ["Score Environnemental (SEC)", f"{summary.get('score_sec', 0)} / 100", _level(summary.get('score_sec', 0))],
        ["Score Social (SSC)", f"{summary.get('score_ssc', 0)} / 100", _level(summary.get('score_ssc', 0))],
        ["Score Souveraineté (SOS)", f"{summary.get('score_sos', 0)} / 100", _level(summary.get('score_sos', 0))],
        ["Score Global", f"{summary.get('score_global', 0)} / 100", _level(summary.get('score_global', 0))],
    ]
    for i, r in enumerate(rows, 1):
        for j, v in enumerate(r):
            table.rows[i].cells[j].text = str(v)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Violations critiques : {counters.get('violations_critiques', 0)} · ").bold = True
    p.add_run(f"Violations graves : {counters.get('violations_graves', 0)} · ").bold = True
    p.add_run(f"Clauses abusives : {counters.get('clauses_abusives', 0)} · ").bold = True
    p.add_run(f"Violations droit national : {counters.get('violations_droit_national', 0)}").bold = True

    # Sections juridique
    if preset in ("juridique", "rejd", "parlementaire"):
        doc.add_page_break()
        doc.add_heading("2. Constats juridiques", level=1)
        jur = report_data.get("juridique") or {}
        for label, key in [
            ("2.1 Violations du droit international", "violations_droit_international"),
            ("2.2 Violations du droit national", "violations_droit_national"),
            ("2.3 Clauses abusives", "clauses_abusives"),
        ]:
            doc.add_heading(label, level=2)
            items = jur.get(key) or []
            if not items:
                doc.add_paragraph("Aucun constat enregistré.")
                continue
            for it in items[:10]:
                p = doc.add_paragraph()
                p.add_run(f"[{it.get('norme_violee') or it.get('code_national_viole') or it.get('type_abus','?')}] ").bold = True
                p.add_run(it.get('nature_violation') or it.get('analyse', '')[:300])
                gp = doc.add_paragraph()
                gr = gp.add_run(f"Gravité : {(it.get('gravite') or '').upper()}")
                gr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B); gr.bold = True
                if it.get("solution"):
                    sp = doc.add_paragraph()
                    sp.add_run("Solution : ").italic = True
                    sp.add_run(it["solution"])

    # Diagnostics
    if preset in ("juridique", "rejd"):
        diag = (report_data.get("diagnostics") or {}).get("fiches", [])
        if diag:
            doc.add_page_break()
            doc.add_heading("3. Fiches diagnostic & moyens de dénonciation", level=1)
            for i, f in enumerate(diag, 1):
                doc.add_heading(f"Fiche {i} — {f.get('anomalie','')[:120]}", level=2)
                p = doc.add_paragraph()
                p.add_run(f"Gravité : {(f.get('gravite') or '').upper()} · Priorité : {(f.get('priorite') or '').upper()}").bold = True
                doc.add_paragraph(f"Qualification juridique : {f.get('qualification_juridique','')}")
                doc.add_paragraph(f"Argument jurisprudentiel : {f.get('argument_jurisprudentiel_principal','')}")
                if f.get("ratio_decidendi"):
                    doc.add_paragraph(f"Ratio decidendi : {f.get('ratio_decidendi','')}")
                if f.get("impact_financier_usd"):
                    doc.add_paragraph(f"Impact financier : {float(f.get('impact_financier_usd') or 0):,.0f} USD")

                if f.get("solutions"):
                    doc.add_heading("Solutions", level=3)
                    for s in f["solutions"][:6]:
                        doc.add_paragraph(f"• {s.get('type','').replace('_',' ').capitalize()} — {s.get('description','')}", style="List Bullet")

                if f.get("moyens_denonciation"):
                    doc.add_heading("Moyens de dénonciation", level=3)
                    for m in f["moyens_denonciation"][:8]:
                        doc.add_paragraph(f"• {m.get('type','').replace('_',' ').capitalize()} — {m.get('description','')} ({m.get('autorite_competente','')})",
                                         style="List Bullet")

    # REJD-specific extended sections
    if preset == "rejd":
        doc.add_page_break()
        doc.add_heading("4. Cadre normatif mobilisé (annexe)", level=1)
        doc.add_paragraph("L'expertise s'appuie sur le corpus normatif suivant :")
        for fam in [
            "Famille 1 — Droit international des ressources naturelles (Rés. 1803, Ruggie, CNUDM, IFC, ITIE...)",
            "Famille 2 — Droit régional africain (Charte africaine Art. 21, Vision minière UA, OHADA)",
            "Famille 4 — Standards contractuels (PSA AIPN, ResourceContracts.org, CNUCED)",
            "Famille 6 — Doctrine juridique (Pacta sunt servanda vs Rebus sic stantibus, contrats léonins)",
        ]:
            doc.add_paragraph(fam, style="List Bullet")

        doc.add_heading("5. Conclusions de l'expertise", level=1)
        niv = (summary.get("niveau_global") or "attention").upper()
        p = doc.add_paragraph()
        p.add_run(f"Au terme de l'analyse, le niveau global de conformité est qualifié de {niv}. ").bold = True
        p.add_run(f"L'expertise relève {counters.get('violations_critiques',0)} violation(s) critique(s), "
                  f"{counters.get('violations_graves',0)} violation(s) grave(s), "
                  f"{counters.get('clauses_abusives',0)} clause(s) abusive(s).")
        doc.add_paragraph(
            "Il est conclu, sous réserve de l'analyse complémentaire d'un avocat qualifié, "
            "à la nécessité de renégocier ou de contester les clauses identifiées, "
            "selon les voies recommandées dans le présent rapport."
        )
        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature.add_run(f"\n\n_____________________\nMéthodologie : Ahmed ELY Mustapha\nRESOURCES-ANALYZER PRO\n{datetime.now().strftime('%d/%m/%Y')}").italic = True

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _level(score):
    s = float(score or 0)
    if s >= 80: return "Conforme"
    if s >= 60: return "Attention"
    if s >= 40: return "Grave"
    return "Critique"


# ============== EXCEL ==============
def generate_excel(project: Dict[str, Any], report_data: Dict[str, Any]) -> bytes:
    wb = Workbook()

    # Sheet 1 — Synthèse
    ws = wb.active; ws.title = "Synthèse"
    ws["A1"] = "RESOURCES-ANALYZER PRO"
    ws["A1"].font = Font(name="Calibri", size=18, bold=True, color=PRIMARY_HEX)
    ws.merge_cells("A1:D1")
    ws["A2"] = f"Projet : {project.get('name', '')}"; ws["A2"].font = Font(bold=True, size=12)
    ws["A3"] = f"Pays : {project.get('country', '')} · Secteur : {project.get('sector', '')} · Ressource : {project.get('resource_type', '')}"
    ws["A4"] = f"Date d'analyse : {datetime.now().strftime('%d/%m/%Y')}"

    summary = report_data.get("summary", {})
    counters = summary.get("compteurs", {})
    fin = summary.get("indicateurs_financiers", {})

    headers = ["Indicateur", "Valeur", "Niveau"]
    for i, h in enumerate(headers, 1):
        c = ws.cell(row=6, column=i, value=h)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor=PRIMARY_HEX)
    rows = [
        ["Score Conformité Juridique", summary.get("score_juridique", 0), _level(summary.get("score_juridique", 0))],
        ["Score Environnemental (SEC)", summary.get("score_sec", 0), _level(summary.get("score_sec", 0))],
        ["Score Social (SSC)", summary.get("score_ssc", 0), _level(summary.get("score_ssc", 0))],
        ["Score Souveraineté (SOS)", summary.get("score_sos", 0), _level(summary.get("score_sos", 0))],
        ["Score Global", summary.get("score_global", 0), _level(summary.get("score_global", 0))],
        ["", "", ""],
        ["Part de l'État (%)", fin.get("part_etat_pct", 0), ""],
        ["Manque à gagner / an (USD)", fin.get("manque_a_gagner_an", 0), ""],
        ["Manque à gagner total (USD)", fin.get("manque_a_gagner_total", 0), ""],
        ["IDC", fin.get("idc", 0), ""],
        ["Cadeau fiscal", str(fin.get("cadeau_fiscal", False)), ""],
        ["", "", ""],
        ["Violations critiques", counters.get("violations_critiques", 0), ""],
        ["Violations graves", counters.get("violations_graves", 0), ""],
        ["Clauses abusives", counters.get("clauses_abusives", 0), ""],
        ["Violations droit national", counters.get("violations_droit_national", 0), ""],
    ]
    for ri, r in enumerate(rows, 7):
        for ci, v in enumerate(r, 1):
            ws.cell(row=ri, column=ci, value=v)
    ws.column_dimensions["A"].width = 38
    ws.column_dimensions["B"].width = 22
    ws.column_dimensions["C"].width = 16

    # Sheet 2 — Violations
    jur = report_data.get("juridique") or {}
    ws2 = wb.create_sheet("Violations")
    h = ["Type", "Norme/Article", "Nature", "Gravité", "Solution"]
    for i, c in enumerate(h, 1):
        cell = ws2.cell(row=1, column=i, value=c)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=PRIMARY_HEX)
    row = 2
    for v in jur.get("violations_droit_international", []):
        ws2.cell(row=row, column=1, value="Droit international")
        ws2.cell(row=row, column=2, value=f"{v.get('norme_violee','')} {v.get('norme_libelle','')}")
        ws2.cell(row=row, column=3, value=v.get("nature_violation",""))
        ws2.cell(row=row, column=4, value=v.get("gravite",""))
        ws2.cell(row=row, column=5, value=v.get("solution",""))
        row += 1
    for v in jur.get("violations_droit_national", []):
        ws2.cell(row=row, column=1, value="Droit national")
        ws2.cell(row=row, column=2, value=f"{v.get('code_national_viole','')} {v.get('article_exact','')}")
        ws2.cell(row=row, column=3, value=v.get("nature_violation",""))
        ws2.cell(row=row, column=4, value=v.get("gravite",""))
        ws2.cell(row=row, column=5, value=v.get("solution",""))
        row += 1
    for c in jur.get("clauses_abusives", []):
        ws2.cell(row=row, column=1, value="Clause abusive")
        ws2.cell(row=row, column=2, value=c.get("type_abus",""))
        ws2.cell(row=row, column=3, value=c.get("analyse",""))
        ws2.cell(row=row, column=4, value=c.get("gravite",""))
        ws2.cell(row=row, column=5, value=c.get("solution",""))
        row += 1
    for col in "ABCDE":
        ws2.column_dimensions[col].width = 30

    # Sheet 3 — Diagnostics
    diag = (report_data.get("diagnostics") or {}).get("fiches", [])
    ws3 = wb.create_sheet("Diagnostics")
    head = ["#", "Anomalie", "Gravité", "Priorité", "Qualification", "Impact financier (USD)", "Argument jurisprudentiel"]
    for i, c in enumerate(head, 1):
        cell = ws3.cell(row=1, column=i, value=c)
        cell.font = Font(bold=True, color="FFFFFF"); cell.fill = PatternFill("solid", fgColor=PRIMARY_HEX)
    for i, f in enumerate(diag, 1):
        ws3.cell(row=i+1, column=1, value=i)
        ws3.cell(row=i+1, column=2, value=f.get("anomalie",""))
        ws3.cell(row=i+1, column=3, value=f.get("gravite",""))
        ws3.cell(row=i+1, column=4, value=f.get("priorite",""))
        ws3.cell(row=i+1, column=5, value=f.get("qualification_juridique",""))
        ws3.cell(row=i+1, column=6, value=f.get("impact_financier_usd",0))
        ws3.cell(row=i+1, column=7, value=f.get("argument_jurisprudentiel_principal",""))
    for col in "ABCDEFG":
        ws3.column_dimensions[col].width = 25

    # Sheet 4 — Données financières
    finr = report_data.get("financier") or {}
    ws4 = wb.create_sheet("Financier")
    items = [
        ["Valeur totale gisement (USD)", finr.get("valeur_totale_gisement", 0)],
        ["Revenus annuels (USD)", finr.get("revenus_annuels", 0)],
        ["Recettes État annuelles (USD)", finr.get("recettes_etat_annuelles", 0)],
        ["Recettes État totales (USD)", finr.get("recettes_etat_totales", 0)],
        ["Part de l'État (%)", finr.get("part_etat_pct", 0)],
        ["Royalties (%)", finr.get("royalties_taux", 0)],
        ["Benchmark Royalties Min (%)", finr.get("royalties_benchmark_min", 0)],
        ["Benchmark Royalties Max (%)", finr.get("royalties_benchmark_max", 0)],
        ["Manque à gagner annuel (USD)", finr.get("manque_a_gagner_annuel", 0)],
        ["Manque à gagner total (USD)", finr.get("manque_a_gagner_total", 0)],
        ["Élément don fiscal", finr.get("element_don_fiscal", 0)],
        ["Cadeau fiscal détecté", str(finr.get("cadeau_fiscal", False))],
    ]
    for i, (k, v) in enumerate(items, 1):
        ws4.cell(row=i, column=1, value=k).font = Font(bold=True)
        ws4.cell(row=i, column=2, value=v)
    ws4.column_dimensions["A"].width = 38
    ws4.column_dimensions["B"].width = 22

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


# ============== ZIP PACK ==============
def generate_rejd_zip(project: Dict[str, Any], report_data: Dict[str, Any], pdf_bytes: bytes) -> bytes:
    word_bytes = generate_word(project, report_data, "rejd")
    excel_bytes = generate_excel(project, report_data)
    out = io.BytesIO()
    safe_name = (project.get("name") or "rapport").replace(" ", "_")[:40]
    date = datetime.now().strftime("%Y%m%d")
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr(f"REJD_{safe_name}_{date}.pdf", pdf_bytes)
        z.writestr(f"REJD_{safe_name}_{date}_editable.docx", word_bytes)
        z.writestr(f"REJD_{safe_name}_{date}_donnees.xlsx", excel_bytes)
        z.writestr(f"REJD_{safe_name}_{date}_donnees_brutes.json",
                   json.dumps(report_data, ensure_ascii=False, indent=2, default=str))
        readme = f"""RAPPORT D'EXPERTISE JURIDIQUE DÉFENDABLE (REJD)
Projet : {project.get('name','')}
Pays : {project.get('country','')}
Secteur : {project.get('sector','')}
Ressource : {project.get('resource_type','')}
Date : {datetime.now().strftime('%d/%m/%Y')}

Contenu du pack :
  1. REJD_{safe_name}_{date}.pdf — Rapport complet (lecture)
  2. REJD_{safe_name}_{date}_editable.docx — Version Word éditable par votre avocat
  3. REJD_{safe_name}_{date}_donnees.xlsx — Données brutes en tableaux multi-onglets
  4. REJD_{safe_name}_{date}_donnees_brutes.json — Toutes les données structurées

Méthodologie : Ahmed ELY Mustapha — Juriste, Expert en Finances Publiques, PMP I-PMP IBM Full Stack Developer.

DOCUMENT PÉDAGOGIQUE — Cette analyse n'a aucune valeur juridictionnelle.
Consultez un avocat qualifié avant toute action.
"""
        z.writestr("README.txt", readme)
    return out.getvalue()
