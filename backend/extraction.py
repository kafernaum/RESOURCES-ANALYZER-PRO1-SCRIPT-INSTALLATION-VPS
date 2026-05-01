"""Document extraction pipeline (PDF / DOCX / XLSX / CSV / TXT) — ZERO LLM here.
LLM is called separately in llm_service.py.
"""
import io
from typing import Tuple


def extract_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            try:
                text_parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(text_parts)
    except Exception as e:
        return f"[Erreur extraction PDF: {e}]"


def extract_docx(file_bytes: bytes) -> str:
    from docx import Document as DocxDocument
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Tables
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[Erreur extraction DOCX: {e}]"


def extract_xlsx(file_bytes: bytes) -> str:
    from openpyxl import load_workbook
    try:
        wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
        out = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            out.append(f"=== Feuille: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                vals = [str(v) for v in row if v is not None]
                if vals:
                    out.append(" | ".join(vals))
        return "\n".join(out)
    except Exception as e:
        return f"[Erreur extraction XLSX: {e}]"


def extract_text(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        return f"[Erreur extraction TXT: {e}]"


def extract_csv(file_bytes: bytes) -> str:
    return extract_text(file_bytes)


def extract_document(filename: str, file_bytes: bytes) -> Tuple[str, str]:
    """Returns (raw_text, detected_extension)."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_pdf(file_bytes), "pdf"
    if name.endswith(".docx") or name.endswith(".doc"):
        return extract_docx(file_bytes), "docx"
    if name.endswith(".xlsx") or name.endswith(".xls"):
        return extract_xlsx(file_bytes), "xlsx"
    if name.endswith(".csv"):
        return extract_csv(file_bytes), "csv"
    if name.endswith(".txt"):
        return extract_text(file_bytes), "txt"
    # fallback
    return extract_text(file_bytes), "txt"
